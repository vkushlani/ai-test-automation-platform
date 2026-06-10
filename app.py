import os
import time
from pathlib import Path

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document as DocxDocument

env_path = Path(__file__).with_name(".env")
load_dotenv(dotenv_path=env_path)

try:
    if "OPENAI_API_KEY" in st.secrets:
        os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
except Exception:
    pass

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

from app_router import classify_question
from app_agents import coordinator_agent
from app_memory_manager import (
    save_memory,
    load_memories,
    save_memory_to_vector_db,
    retrieve_memory,
)

# =====================================================
# PAGE CONFIG
# =====================================================

st.set_page_config(
    page_title="AI-Powered QA Agent Platform",
    layout="wide"
)

st.title("AI-Powered QA Agent Platform")

st.caption(
    "AI-powered Testing & Automation Assistant for test design, Selenium framework generation, defect analysis, traceability, regression risk assessment, document intelligence, and QA planning."
)

with st.expander("🚀 Platform Capabilities", expanded=True):
    st.markdown("""
### 🤖 AI Testing & Automation Agents

✅ AI-Powered Test Case Generation  
✅ Selenium + Cucumber + TestNG Framework Generation  
✅ Automation Framework Design (POM)  
✅ Defect & Root Cause Analysis  
✅ Requirement Traceability Support  
✅ Regression Risk Assessment  
✅ Test Coverage Analysis  

---

### 📄 Document Intelligence

📄 Summarize Uploaded Documents  
📄 Compare Multiple Documents  
📄 Analyze Release Notes  
📄 Generate Tests from Requirements & User Stories  

---

### 🛠 Automation Engineering

⚙️ Generate Feature Files  
⚙️ Generate Step Definitions  
⚙️ Generate Page Objects  
⚙️ Generate Maven Project Structure  
⚙️ Download Automation Framework ZIP  

---

### 🚀 Advanced AI Features

🧠 Persistent Memory  
🔎 RAG Search  
🌐 Live Web Search  
🤖 Multi-Agent Architecture  
📂 Multi-Document Analysis  
💬 Conversational QA Assistant  
""")

# =====================================================
# LLM
# =====================================================

llm = ChatOpenAI(
    model="gpt-4.1-mini",
    temperature=0.2
)

# =====================================================
# SESSION STATE
# =====================================================

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "framework_zip_path" not in st.session_state:
    st.session_state.framework_zip_path = None

if "uploaded_file_names" not in st.session_state:
    st.session_state.uploaded_file_names = []

if "uploaded_file_contents" not in st.session_state:
    st.session_state.uploaded_file_contents = {}

past_memories = load_memories()
DEBUG_MODE = False

# =====================================================
# FILE EXTRACTION
# =====================================================

def extract_file_content(uploaded_file):
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if file_name.endswith(".pdf"):
        pdf_reader = PdfReader(uploaded_file)
        text = ""

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        return text

    if file_name.endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text

    if file_name.endswith(".csv"):
        encodings = ["utf-8", "cp1252", "latin-1", "ISO-8859-1"]

        for encoding in encodings:
            try:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding=encoding)
                return df.to_string(index=False)
            except Exception:
                continue

        return "Unable to read CSV file."

    if file_name.endswith(".xlsx"):
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file)
        return df.to_string(index=False)

    return ""

# =====================================================
# OPTIONAL DOCUMENT UPLOAD
# =====================================================

uploaded_files = st.file_uploader(
    "＋ Attach documents",
    type=["txt", "pdf", "docx", "csv", "xlsx"],
    accept_multiple_files=True,
    help="Optional: attach QA documents such as requirements, release notes, defect reports, test plans, CSV, Excel, Word, or PDF files."
)

if uploaded_files:
    st.session_state.uploaded_file_names = []
    st.session_state.uploaded_file_contents = {}

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_content = extract_file_content(uploaded_file)

        st.session_state.uploaded_file_names.append(file_name)
        st.session_state.uploaded_file_contents[file_name] = file_content

    uploaded_names = ", ".join(st.session_state.uploaded_file_names)

    st.success(
        f"✅ {len(st.session_state.uploaded_file_names)} document(s) uploaded: {uploaded_names}"
    )

# =====================================================
# CHAT INPUT
# =====================================================

question = st.chat_input("Ask the AI Testing Agent...")

# =====================================================
# DISPLAY CHAT HISTORY
# =====================================================

st.subheader("Conversation")

for message in st.session_state.chat_history:
    if isinstance(message, dict):
        with st.chat_message(message.get("role", "assistant")):
            st.write(message.get("content", ""))

            if message.get("role") == "assistant" and "response_time" in message:
                st.caption(
                    f"⏱️ Response generated in {message['response_time']:.2f} seconds"
                )

# =====================================================
# MAIN PROCESS
# =====================================================

if question:
    start_time = time.time()

    with st.chat_message("user"):
        st.write(question)

    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):

            query_type = classify_question(question)

            if query_type != "automation_framework":
                st.session_state.framework_zip_path = None

            document_required_workflows = [
                "comparison",
                "summary",
                "document_count",
                "document_names",
                "automation_framework",
            ]

            has_uploaded_docs = bool(st.session_state.uploaded_file_names)

            has_chat_context = any(
                isinstance(msg, dict)
                and msg.get("role") == "assistant"
                and (
                    "test case" in msg.get("content", "").lower()
                    or "scenario" in msg.get("content", "").lower()
                    or "expected result" in msg.get("content", "").lower()
                )
                for msg in st.session_state.chat_history
            )

            can_continue = True
            answer = ""

            if query_type in document_required_workflows and not has_uploaded_docs:
                if query_type == "automation_framework" and has_chat_context:
                    can_continue = True
                else:
                    answer = (
                        "This request requires uploaded documents or previous manual test cases. "
                        "Please upload manual test cases, requirements, or user stories, "
                        "or first generate manual test cases, then ask me to convert them into "
                        "a Selenium Cucumber TestNG automation framework."
                    )
                    can_continue = False

            if can_continue:
                all_documents = []
                uploaded_document_names = []
                combined_content = ""
                unique_documents = []
                retrieved_context = ""
                memory_context = ""

                history = "\n".join(
                    [
                        f"{msg.get('role', '')}: {msg.get('content', '')}"
                        for msg in st.session_state.chat_history
                        if isinstance(msg, dict)
                    ]
                )

                # Read uploaded documents
                if st.session_state.uploaded_file_contents:
                    for filename, content in st.session_state.uploaded_file_contents.items():
                        uploaded_document_names.append(filename)

                        combined_content += f"""

DOCUMENT NAME:
{filename}

DOCUMENT CONTENT:
{content}

"""

                        all_documents.append(
                            Document(
                                page_content=content,
                                metadata={"source": filename}
                            )
                        )

                    unique_documents = list(set(uploaded_document_names))

                # Disable memory when uploaded docs are present
                if query_type in [
                    "document_count",
                    "document_names",
                    "comparison",
                    "summary",
                    "automation_framework",
                ]:
                    memory_results = []
                else:
                    if st.session_state.uploaded_file_names:
                        memory_results = []
                    else:
                        memory_results = retrieve_memory(question)

                for memory in memory_results:
                    memory_context += f"""

{memory.page_content}

"""

                # RAG search only if uploaded documents exist
                if all_documents:
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=80
                    )

                    chunks = splitter.split_documents(all_documents)

                    embedding_model = OpenAIEmbeddings()

                    vectorstore = Chroma.from_documents(
                        documents=chunks,
                        embedding=embedding_model,
                        persist_directory="vector_db"
                    )

                    results = vectorstore.similarity_search(
                        question,
                        k=6
                    )

                    for r in results:
                        source = r.metadata.get("source", "Unknown")

                        retrieved_context += f"""

DOCUMENT:
{source}

CONTENT:
{r.page_content}

"""

                # Master context
                if st.session_state.uploaded_file_names:
                    master_context = f"""

Uploaded Documents:
{unique_documents}

Retrieved Context:
{retrieved_context}

Document Content:
{combined_content}

IMPORTANT:
Use ONLY the uploaded document content.
Ignore previous conversation history.
Ignore historical memory unless explicitly requested.
"""
                else:
                    master_context = f"""

Conversation History:
{history}

Historical Memory:
{memory_context}

Retrieved Context:
{retrieved_context}

Question:
{question}
"""

                # Document count
                if query_type == "document_count":
                    answer = f"""
Total uploaded documents: {len(unique_documents)}

Document names:

{chr(10).join(unique_documents)}
"""

                # Document names
                elif query_type == "document_names":
                    answer = f"""
Uploaded document name(s):

{chr(10).join(unique_documents)}
"""

                # Summary
                elif query_type == "summary":
                    prompt = f"""
You are a document summarization expert.

Summarize all uploaded documents clearly.

Documents:
{unique_documents}

Document Content:
{combined_content}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

                # Comparison
                elif query_type == "comparison":
                    prompt = f"""
You are a document comparison expert.

Compare all uploaded documents.

Include:
1. Similarities
2. Differences
3. Key findings
4. Missing or unique topics

Documents:
{unique_documents}

Document Content:
{combined_content}

Question:
{question}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

                # Multi-agent workflows
                elif query_type in [
                    "test_case",
                    "defect_analysis",
                    "traceability",
                    "regression_risk",
                    "coverage_pipeline",
                    "automation",
                    "automation_framework",
                    "website_testing",
                    "planning",
                    "release_readiness",
                    "web_search",
                ]:
                    answer = coordinator_agent(
                        query_type=query_type,
                        context=master_context,
                        question=question
                    )

                    if query_type == "automation_framework":
                        st.session_state.framework_zip_path = (
                            "exports/generated_selenium_cucumber_framework.zip"
                        )

                # Default QA / RAG workflow
                else:
                    if retrieved_context:
                        prompt = f"""
You are a Senior QA AI Assistant.

Use uploaded documents if relevant.

Retrieved Context:
{retrieved_context}

Question:
{question}
"""
                    else:
                        prompt = f"""
You are a friendly AI QA Assistant.

If the user is chatting casually, respond naturally.
If the user asks QA or software testing questions, answer using your QA knowledge.

Conversation History:
{history}

Historical Memory:
{memory_context}

Question:
{question}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

            end_time = time.time()
            response_time = end_time - start_time

            st.write(answer)

            st.caption(
                f"⏱️ Response generated in {response_time:.2f} seconds"
            )

            st.session_state.chat_history.append(
                {
                    "role": "user",
                    "content": question,
                }
            )

            st.session_state.chat_history.append(
                {
                    "role": "assistant",
                    "content": answer,
                    "response_time": response_time,
                }
            )

            save_memory(question, answer)

            save_memory_to_vector_db(question, answer)

            if DEBUG_MODE:
                with st.expander("Retrieved Document Context"):
                    st.text(retrieved_context if "retrieved_context" in locals() else "")

                with st.expander("Retrieved Memory Context"):
                    st.text(memory_context if "memory_context" in locals() else "")

                with st.expander("Persistent Memory Store"):
                    st.write(past_memories)

# =====================================================
# DOWNLOAD GENERATED FRAMEWORK
# =====================================================

if (
    st.session_state.framework_zip_path
    and os.path.exists(st.session_state.framework_zip_path)
):
    with open(st.session_state.framework_zip_path, "rb") as zip_file:
        st.download_button(
            label="⬇️ Download Selenium Cucumber Framework ZIP",
            data=zip_file,
            file_name="generated_selenium_cucumber_framework.zip",
            mime="application/zip",
            key="download_framework_zip"
        )

# =====================================================
# FOOTER
# =====================================================

st.markdown("---")

st.caption(
    """
   <b>Built by Vikas Kushlani</b><br>
   
    Streamlit • LangChain • OpenAI • ChromaDB • Multi-Agent AI
    """,
    unsafe_allow_html=True
)